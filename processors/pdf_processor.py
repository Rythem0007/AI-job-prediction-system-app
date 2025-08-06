import io
import re
from typing import Optional, Union
import streamlit as st

class PDFProcessor:
    def __init__(self):
        pass
    
    def extract_text(self, file_upload) -> str:
        """Extract text from uploaded PDF or DOC file"""
        if file_upload is None:
            return ""
        
        file_extension = file_upload.name.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_upload)
            elif file_extension in ['doc', 'docx']:
                return self._extract_from_doc(file_upload)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_extension.upper()}: {str(e)}")
    
    def _extract_from_pdf(self, file_upload) -> str:
        """Extract text from PDF file"""
        try:
            # Try PyPDF2 first
            import PyPDF2
            
            # Read file content
            file_content = file_upload.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if text.strip():
                return self._clean_extracted_text(text)
            
        except ImportError:
            pass  # PyPDF2 not available
        except Exception as e:
            st.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        try:
            # Try pdfplumber as fallback
            import pdfplumber
            
            file_upload.seek(0)  # Reset file pointer
            file_content = file_upload.read()
            
            text = ""
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return self._clean_extracted_text(text)
                
        except ImportError:
            pass  # pdfplumber not available
        except Exception as e:
            st.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # If all PDF extraction methods fail, try to read as text
        try:
            file_upload.seek(0)
            content = file_upload.read()
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        text = content.decode(encoding)
                        return self._clean_extracted_text(text)
                    except UnicodeDecodeError:
                        continue
        except Exception as e:
            st.warning(f"Text extraction as fallback failed: {str(e)}")
        
        raise Exception("Unable to extract text from PDF. Please ensure the file is not password protected or corrupted.")
    
    def _extract_from_doc(self, file_upload) -> str:
        """Extract text from DOC/DOCX file"""
        file_extension = file_upload.name.lower().split('.')[-1]
        
        try:
            # Try python-docx for DOCX files
            if file_extension == 'docx':
                try:
                    import docx
                except ImportError:
                    raise Exception("python-docx not installed. Cannot process DOCX files.")
                
                file_content = file_upload.read()
                doc = docx.Document(io.BytesIO(file_content))
                
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
                
                if text.strip():
                    return self._clean_extracted_text(text)
                    
        except ImportError:
            st.warning("python-docx not available. Cannot process DOCX files.")
        except Exception as e:
            st.warning(f"DOCX extraction failed: {str(e)}")
        
        # For DOC files or if DOCX extraction fails, try alternative methods
        try:
            # Try reading as plain text (sometimes works for simple DOC files)
            file_upload.seek(0)
            content = file_upload.read()
            
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        text = content.decode(encoding, errors='ignore')
                        # Remove common binary artifacts
                        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII
                        text = re.sub(r'\x00+', ' ', text)  # Remove null bytes
                        
                        if len(text.strip()) > 50:  # Ensure we got meaningful text
                            return self._clean_extracted_text(text)
                    except Exception:
                        continue
                        
        except Exception as e:
            st.warning(f"Alternative DOC extraction failed: {str(e)}")
        
        raise Exception(f"Unable to extract text from {file_extension.upper()} file. Please convert to PDF or ensure the file is not corrupted.")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces
        
        # Remove common OCR artifacts
        text = re.sub(r'[^\w\s\-\.\,\(\)\[\]\{\}\:\/\@\+\#\&\%\$\!]', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Clean up common formatting issues
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)  # Fix word breaks across lines
        
        return text.strip()
    
    def validate_file(self, file_upload) -> tuple[bool, str]:
        """Validate uploaded file"""
        if file_upload is None:
            return False, "No file uploaded"
        
        # Check file size (limit to 10MB)
        file_size = len(file_upload.getvalue())
        if file_size > 10 * 1024 * 1024:  # 10MB limit
            return False, "File size too large. Please upload a file smaller than 10MB."
        
        # Check file extension
        allowed_extensions = ['pdf', 'doc', 'docx']
        file_extension = file_upload.name.lower().split('.')[-1]
        
        if file_extension not in allowed_extensions:
            return False, f"Unsupported file format. Please upload: {', '.join(allowed_extensions.upper())}"
        
        return True, "File is valid"
    
    def get_file_info(self, file_upload) -> dict:
        """Get information about the uploaded file"""
        if file_upload is None:
            return {}
        
        file_size = len(file_upload.getvalue())
        file_extension = file_upload.name.lower().split('.')[-1]
        
        return {
            'filename': file_upload.name,
            'size': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'extension': file_extension,
            'mime_type': file_upload.type if hasattr(file_upload, 'type') else 'unknown'
        }
